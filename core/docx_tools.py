import re
from pathlib import Path
from typing import Optional

import docx
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def open_docx(path: str | Path) -> Document:
    """
    Открывает существующий документ Word.
    :param path: Путь к файлу docx
    :return: Объект Document
    """
    path = Path(path)
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f'Файл {path} не найден')
    return Document(path)


def _extract_after_label(doc: Document, label: str) -> Optional[str]:
    """
    Поиск метки (label) и возвращение текста следующего абзаца после неё
    """

    for i, para in enumerate(doc.paragraphs):
        if label in para.text:
            if i + 1 < len(doc.paragraphs):
                return doc.paragraphs[i + 1].text.strip()
            else:
                return None
    return None


def extract_fio_debtor(doc: Document) -> Optional[str]:
    """
    Извлечение ФИО должника. Метка: 'Должник:'
    """
    return _extract_after_label(doc, 'Должник:')


# def extract_fio_financial_manager(docx_path: Path) -> Optional[str]:
#     """
#     Извлечение ФИО фин. управляющего. Метка: 'Финансовый управляющий:'
#     """
#     doc = _open_docx(docx_path)
#     return _extract_after_label(doc, 'Финансовый управляющий:')


def extract_case_number(doc: Document) -> Optional[str]:
    """
    Извлечение номера дела. Формат: Дело № A33-12345/2024
    """
    pattern = r'Дело\s+№\s*([AА]\d{2}-\d+/\d{4})'

    for para in doc.paragraphs:
        match = re.search(pattern, para.text)
        if match:
            return match.group(1)

    return None


def format_appendices(doc: Document) -> None:
    """
    Форматирует блок "ПРИЛОЖЕНИЯ:" в документе как нумерованный список.
    Убирает старые цифры перед созданием формата списка.
    """
    appendices_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == 'ПРИЛОЖЕНИЯ:' and any(run.bold for run in para.runs):
            appendices_start = i + 1
            break

    if appendices_start is None:
        return  # Блок не найден

    num_id = 1  # id нумерации Word, можно оставить 1 для простого случая

    for para in doc.paragraphs[appendices_start:]:
        text = para.text.strip()
        if text and re.match(r'^\d+\.\s+', text):
            run_formats = [(run.font.name, run.font.size, run.font.bold, run.font.italic) for run in para.runs]

            para.text = re.sub(r'^\d+\.\s+', '', text)

            p = para._p
            pPr = p.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')

            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numPr.append(ilvl)

            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), str(num_id))
            numPr.append(numId)

            pPr.append(numPr)

            for run, (name, size, bold, italic) in zip(para.runs, run_formats):
                run.font.name = name
                run.font.size = size
                run.font.bold = bold
                run.font.italic = italic
        else:
            break


def insert_bank_table(doc_statement: Document, doc_requisities: Document, bank_name: str):
    """
    Вставляет таблицу с реквизитами нужного банка в заявление.

    :param statement_path: Путь к заявлению
    :param path_banks_file: Путь к файлу с реквизитами банков
    :param bank_name: Название банка, чьи реквизиты нужно вставить
    """
    # --- Находим таблицу нужного банка ---
    target_table: Table | None = None
    for idx, para in enumerate(doc_requisities.paragraphs):
        if bank_name.lower() in para.text.lower():
            # ищем первую таблицу после параграфа
            for child in doc_requisities.element.body[idx + 1 :]:
                if child.tag.endswith('tbl'):
                    target_table = Table(child, doc_requisities)
                    break
            break

    if target_table is None:
        raise ValueError(f"Таблица с реквизитами '{bank_name}' не найдена")

    # --- Находим таблицу в заявлении ---
    stmt_table: Table | None = None
    for idx, para in enumerate(doc_statement.paragraphs):
        if para.text.startswith('Реквизиты ПАО Сбербанк для погашения задолженности'):
            # ищем первую таблицу после параграфа
            for child in doc_statement.element.body[idx + 1 :]:
                if child.tag.endswith('tbl'):
                    stmt_table = Table(child, doc_statement)
                    break
            break

    if stmt_table is None:
        raise ValueError('Таблица для вставки реквизитов в заявлении не найдена')

    # --- Проверка размеров таблиц ---
    if len(stmt_table.rows) != len(target_table.rows) or len(stmt_table.columns) != len(target_table.columns):
        raise ValueError('Таблицы имеют разную размерность')

    # --- Копирование содержимого с сохранением стиля ---
    for i, row in enumerate(target_table.rows):
        for j, cell in enumerate(row.cells):
            stmt_cell = stmt_table.cell(i, j)
            stmt_cell.text = ''
            for run in cell.paragraphs[0].runs:
                new_run = stmt_cell.paragraphs[0].add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb


def get_bank_list(doc: Document) -> list[str]:
    """Извлекает список банков из документа."""
    # BANK_FILE = Path('settings/bank_requisites.docx')
    banks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        # Если абзац непустой и перед ним идет таблица (или просто проверяем на слово "Реквизиты")
        if text.lower().startswith('реквизиты'):
            # Берем только название банка, удаляя слово "Реквизиты"
            bank_name = text.split('Реквизиты')[-1].strip()
            if bank_name:
                banks.append(bank_name)
    return banks


def delete_words(doc: Document, targets: list[str]) -> None:
    """
    Удаляет слова/фразы из документа, сбрасывая форматирование,
    но сохраняя размер шрифта каждого абзаца.
    """
    for para in doc.paragraphs:
        if any(target in para.text for target in targets):
            # Сохраняем размер шрифта из первого run, если есть
            font_size = None
            if para.runs:
                font_size = para.runs[0].font.size

            # Удаляем слова/фразы
            text = para.text
            for target in targets:
                text = text.replace(target, '')
            para.text = text

            # Применяем сохранённый размер шрифта ко всему абзацу
            if font_size:
                for run in para.runs:
                    run.font.size = font_size


def delete_paragraphs(doc: Document, targets: list[str]) -> None:
    """
    Удаляет целиком абзацы, в которых встречаются слова/фразы из targets.
    """
    # Создаём копию списка абзацев, чтобы безопасно удалять во время итерации
    for para in list(doc.paragraphs):
        if any(target in para.text for target in targets):
            # Удаляем параграф через XML
            p_element = para._element
            p_element.getparent().remove(p_element)


def insert_gosposhlina(doc: Document, template: Document):
    found_prosit = False
    inserted = False

    for idx, para in enumerate(doc.paragraphs):
        if not found_prosit:
            if para.text.strip() == 'ПРОСИТ СУД:':
                found_prosit = True
            continue
        else:
            # Находим параграф с "1." чтобы взять ФИО
            if para.text.strip().startswith('1.'):
                text_1 = para.text
                # Ищем ФИО между "кредиторов" и "в размере"
                start = text_1.find('кредиторов')
                end = text_1.find('в размере')
                if start != -1 and end != -1:
                    fio_debtor = text_1[start + len('кредиторов') : end].strip()
                else:
                    fio_debtor = 'ФИО'

            if not inserted and para.text.strip().startswith('2.'):
                para_to_copy = template.paragraphs[0]
                # Заменяем <ФИО> на найденное ФИО прямо в XML
                p_xml = para_to_copy._p.xml.replace('ФИО', fio_debtor)
                # Вставляем новый параграф с уже заменённым текстом
                new_p_element = para._element.addprevious(parse_xml(p_xml))
                Paragraph(new_p_element, para._parent)
                inserted = True
                start_idx = idx + 1  # Продолжаем с текущего индекса для изменения нумерации
                break

    # Если вставили, идём дальше и увеличиваем номера
    if inserted:
        for para in doc.paragraphs[start_idx:]:
            text = para.text.strip()
            if text.startswith('ПРИЛОЖЕНИЯ'):
                break
            # Проверяем, начинается ли параграф с числа и точки
            if len(text) > 1 and text[0].isdigit() and text[1] == '.':
                num = int(text[0]) + 1
                # Изменяем первый run, если он содержит цифру
                for run in para.runs:
                    if run.text.startswith(text[0] + '.'):
                        # Сохраняем стиль run, меняем только текст цифры
                        run.text = f'{num}.' + run.text[2:]
                        break


if __name__ == '__main__':
    # Пример использования
    try:
        # path_del_paragraphs = 'templates/del_paragraphs.txt'
        # with open(path_del_paragraphs, 'r', encoding='utf-8') as file:
        #     words = [line.strip('\n') for line in file if line.strip()]
        template = Document(r'templates\gosposhlina.docx')  # Загружаем шаблон
        doc = open_docx('заявление на включение требований в РТК_2rsfdofiswdf.docx.docx')
        insert_gosposhlina(doc, template)
        doc.save('output.docx')

    except Exception as e:
        print(f'Ошибка: {e}')
