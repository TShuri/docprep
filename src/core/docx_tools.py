import re
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


def open_docx(path: str | Path) -> Document:
    """
    Открывает существующий документ Word.
    :param path: Путь к файлу docx
    :return: Объект Document
    """
    path = Path(path)
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f'Файл {path} не найден')
    return Document(path)


def _extract_after_label(doc: Document, label: str) -> Optional[str]:
    """
    Поиск метки (label) и возвращение текста следующего абзаца после неё
    """

    for i, para in enumerate(doc.paragraphs):
        if label in para.text:
            if i + 1 < len(doc.paragraphs):
                return doc.paragraphs[i + 1].text.strip()
            else:
                return None
    return None


def _force_font_size(paragraph, font_name='Times New Roman', size=11):
    """
    Жёстко задаёт Times New Roman для всех runs в параграфе.
    Если runs нет, создаёт один run с пустым текстом.
    """
    if not paragraph.runs:  # если runs пустые
        run = paragraph.add_run('')
        run.font.name = font_name
        run.font.size = Pt(size)
    else:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)


def extract_fio_debtor(doc: Document) -> Optional[str]:
    """
    Извлечение ФИО должника. Метка: 'Должник:'
    """
    return _extract_after_label(doc, 'Должник:')


def extract_fio_financial_manager(doc: Document) -> Optional[str]:
    """
    Извлечение ФИО фин. управляющего. Метка: 'Финансовый управляющий:'
    """
    return _extract_after_label(doc, 'Финансовый управляющий:')


def extract_case_number(doc: Document) -> Optional[str]:
    """
    Извлечение номера дела. Формат: Дело № A33-12345/2024
    """
    pattern = r'Дело\s+№\s*([AА]\d{2}-\d+/\d{4})'

    for para in doc.paragraphs:
        match = re.search(pattern, para.text)
        if match:
            return match.group(1)

    return None


def extract_date(doc: Document) -> Optional[str]:
    """
    Извлекает дату решения из параграфа, который идёт через
    параграфа с текстом 'о включении требований в реестр кредиторов'.
    Формат даты: dd.mm.yyyy
    """
    target_text = 'о включении требований в реестр кредиторов'
    date_pattern = r'\b(\d{2}\.\d{2}\.\d{4})\b'

    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs[:-2]):  # не смотрим последние 2 параграфа
        if target_text in para.text:
            next_para_text = paragraphs[i + 2].text
            match = re.search(date_pattern, next_para_text)
            if match:
                return match.group(1)
    return None


def delete_words_in_obyazatelstvo(doc: Document, targets: list[str]) -> None:
    """
    Удаляет слова/фразы из документа только в диапазоне
    от "Обязательство №" до "ПРОСИТ СУД:".
    При этом сбрасывает форматирование текста, но сохраняет
    размер шрифта и выставляет Times New Roman.
    """
    start_found = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if 'Обязательство №' in text:
            start_found = True
            continue

        if start_found and 'ПРОСИТ СУД:' in text:
            break

        if start_found and any(target in para.text for target in targets):
            new_text = para.text
            for target in targets:
                new_text = new_text.replace(target, '')
            para.text = new_text

            _force_font_size(para)


def delete_paragraphs_in_obyazatelstvo(doc: Document, targets: list[str]) -> None:
    """
    Удаляет целиком абзацы, содержащие слова/фразы из targets,
    но только в диапазоне от абзаца с "Обязательство №"
    до абзаца с "ПРОСИТ СУД:" (не включая его)
    """
    start_found = False

    for para in list(doc.paragraphs):
        text = para.text.strip()

        if 'Обязательство №' in text:
            start_found = True
            continue

        if start_found and 'ПРОСИТ СУД:' in text:
            break

        if start_found and any(target in para.text for target in targets):
            p_element = para._element
            p_element.getparent().remove(p_element)


def delete_paragraphs_in_gosposhlina(doc: Document, targets: list[str]) -> None:
    """Удаляет целиком абзацы, содержащие слова/фразы из targets,
    но только в диапазоне от абзаца с "ПРОСИТ СУД:"
    до абзаца с "ПРИЛОЖЕНИЧ:" (не включая его)"""
    start_found = False

    paragraphs = list(doc.paragraphs)
    for idx, para in enumerate(paragraphs):
        text = para.text.strip()

        if 'ПРОСИТ СУД:' in text:
            start_found = True
            continue

        if start_found and 'ПРИЛОЖЕНИЯ:' in text:
            break

        if start_found and any(target in para.text for target in targets):
            p_element = para._element
            p_element.getparent().remove(p_element)

            if idx > 0:  # удаляем предыдущий абзац(пустая строка), если он существует
                prev_para = paragraphs[idx - 1]
                prev_el = prev_para._element
                if prev_el.getparent() is not None:  # чтобы не упасть, если уже удалён
                    prev_el.getparent().remove(prev_el)


def insert_gosposhlina(doc: Document, template: Document):
    """
    Вставляет в судебную часть вставку про оплату госпошлины.
    Берет ФИО из судебной части.
    Следующие пункты увеличивает на +1.
    """
    start_found = False
    inserted = False
    fio_debtor = ''

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if 'ПРОСИТ СУД:' in text:
            start_found = True
            continue

        if start_found and 'ПРИЛОЖЕНИЯ:' in text:
            break

        if start_found and not inserted:
            if text.startswith('1.'):  # Находим параграф с "1." чтобы взять ФИО
                text_1 = para.text
                start = text_1.find('кредиторов')
                end = text_1.find('в размере')
                if start != -1 and end != -1:  # Ищем ФИО между "кредиторов" и "в размере"
                    fio_debtor = text_1[start + len('кредиторов') : end].strip()
                    continue

            if text and text[0].isdigit() and text[1] == '.':
                # --- вставка 1-го параграфа шаблона (Пункт про Госпошлину) ---
                para_to_copy = template.paragraphs[0]
                p_xml = para_to_copy._p.xml.replace('ФИО', fio_debtor)
                e = parse_xml(p_xml)
                para._element.addprevious(e)
                new_para = Paragraph(e, para._parent)
                _force_font_size(new_para, size=11)

                # --- вставка 2-го параграфа шаблона (абзац с размером Pt5 - пустая строка) ---
                para_to_copy2 = template.paragraphs[1]
                p_xml2 = para_to_copy2._p.xml
                e2 = parse_xml(p_xml2)
                new_para._element.addnext(e2)
                new_para2 = Paragraph(e2, para._parent)
                _force_font_size(new_para2, size=5)

                inserted = True
                start_idx = idx + 1  # Продолжаем с текущего индекса для изменения нумерации

        # Если вставили, идём дальше и увеличиваем номера
        if start_found and inserted:
            current_number = 3
            for para in doc.paragraphs[start_idx:]:
                text = para.text.strip()
                if text.startswith('ПРИЛОЖЕНИЯ'):
                    break

                if len(text) > 1 and text[0].isdigit() and text[1] == '.':
                    # Сохраняем run, удаляя первую цифру с точкой
                    original_runs = []
                    first_run_skipped = False
                    for r in para.runs:
                        r_text = r.text
                        if not first_run_skipped:
                            # Ищем позицию точки после цифры
                            dot_idx = r_text.find('.')
                            if dot_idx != -1:
                                r_text = r_text[dot_idx + 1 :]  # убираем цифру и точку
                                first_run_skipped = True
                            else:
                                # весь run игнорируем, если точка ещё не найдена
                                continue
                        # Создаём новый run XML с обрезанным текстом
                        new_r = deepcopy(r._element)
                        new_r.text = r_text  # устанавливаем обрезанный текст
                        original_runs.append(new_r)

                    # Очищаем абзац
                    para.clear()

                    # Новый run для номера
                    run_num = para.add_run(f'{current_number}.')
                    run_num.font.name = 'Times New Roman'
                    run_num.font.size = Pt(11)
                    run_num.font.bold = True

                    # Вставляем оставшиеся run
                    for r_element in original_runs:
                        para._element.append(deepcopy(r_element))

                    current_number += 1


def delete_paragraphs_in_appendices(doc: Document, targets: list[str]) -> None:
    """
    Удаляет целиком абзацы, содержащие слова/фразы из targets,
    но только в диапазоне от абзаца с "ПРИЛОЖЕНИЯ:"
    до абзаца с "Реквизиты ПАО Сбербанк" (не включая его)
    """
    start_found = False

    for para in list(doc.paragraphs):
        text = para.text.strip()

        if 'ПРИЛОЖЕНИЯ:' in text:
            start_found = True
            continue

        if start_found and 'Реквизиты ПАО Сбербанк' in text:
            break

        if start_found and any(target in para.text for target in targets):
            p_element = para._element
            p_element.getparent().remove(p_element)


def format_appendices(doc: Document) -> None:
    """
    Форматирует блок "ПРИЛОЖЕНИЯ:" в документе как нумерованный список.
    Убирает старые цифры перед созданием формата списка.
    """
    appendices_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().upper() == 'ПРИЛОЖЕНИЯ:' and any(run.bold for run in para.runs):
            appendices_start = i + 1
            break

    if appendices_start is None:
        return  # Блок не найден

    num_id = 1  # id нумерации Word, можно оставить 1 для простого случая

    for para in doc.paragraphs[appendices_start:]:
        text = para.text.strip()
        if text and re.match(r'^\d+\.\s+', text):
            run_formats = [(run.font.name, run.font.size, run.font.bold, run.font.italic) for run in para.runs]

            para.text = re.sub(r'^\d+\.\s+', '', text)

            p = para._p
            pPr = p.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')

            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numPr.append(ilvl)

            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), str(num_id))
            numPr.append(numId)

            pPr.append(numPr)

            for run, (name, size, bold, italic) in zip(para.runs, run_formats):
                run.font.name = name
                run.font.size = size
                run.font.bold = bold
                run.font.italic = italic
        else:
            break


def insert_bank_table(doc_statement: Document, doc_requisities: Document, bank_name: str):
    """
    Вставляет таблицу с реквизитами нужного банка в заявление.

    :param statement_path: Путь к заявлению
    :param path_banks_file: Путь к файлу с реквизитами банков
    :param bank_name: Название банка, чьи реквизиты нужно вставить
    """
    # --- Находим таблицу нужного банка ---
    target_table: Table | None = None
    for table in doc_requisities.tables:
        for row in table.rows:
            for cell in row.cells:
                if bank_name.lower() in cell.text.lower():
                    target_table = table
                    break
            if target_table:
                break
        if target_table:
            break

    if target_table is None:
        raise ValueError(f"Таблица с реквизитами '{bank_name}' не найдена")

    # --- Находим таблицу в заявлении ---
    stmt_table: Table | None = None
    for idx, para in enumerate(doc_statement.paragraphs):
        if para.text.startswith('Реквизиты ПАО Сбербанк для погашения задолженности'):
            for child in doc_statement.element.body[idx + 1 :]:
                if child.tag.endswith('tbl'):  # ищем первую таблицу после параграфа
                    stmt_table = Table(child, doc_statement)
                    break
            break

    if stmt_table is None:
        raise ValueError('Таблица для вставки реквизитов в заявлении не найдена')

    # --- Проверка размеров таблиц ---
    if len(stmt_table.rows) != len(target_table.rows) or len(stmt_table.columns) != len(target_table.columns):
        raise ValueError('Таблицы имеют разную размерность')

    # --- Копирование содержимого с сохранением стиля ---
    for i, row in enumerate(target_table.rows):
        for j, cell in enumerate(row.cells):
            stmt_cell = stmt_table.cell(i, j)
            stmt_cell.text = ''

            src_para = cell.paragraphs[0]
            dst_para = stmt_cell.paragraphs[0]

            for run in src_para.runs:
                dst_para.add_run(run.text)

            _force_font_size(dst_para, size=11)


def get_bank_list(doc: Document) -> list[str]:
    """Извлекает список банков из документа с реквизитами банков."""
    banks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.lower().startswith('реквизиты'):
            bank_name = text.split('Реквизиты')[-1].strip()
            if bank_name:
                banks.append(bank_name)
    return banks


def insert_signature(doc: Document, signa_path: Path):
    """
    Добавляет подпись в самый последний абзац документа и центрирует её.

    :param doc: Document - объект python-docx
    :param signa_path: Path - путь к картинке подписи
    :param width_cm: ширина подписи в сантиметрах (по умолчанию 5 см)
    """
    last_paragraph = doc.paragraphs[-1]  # Берём последний абзац
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Центрируем абзац

    run = last_paragraph.add_run()  # Вставляем подпись
    run.add_picture(str(signa_path), width=Cm(3))


def insert_zalog_contacts(doc: Document, template: Document):
    """
    Вставляет залоговые контакты из шаблона в документ после абзаца,
    который содержит текст 'Электронный адрес: Bankrot_FL@sberbank.ru'
    Если в документе уже есть контакты, то не будет вставлять.
    """
    target_text = 'Электронный адрес: Bankrot_FL@sberbank.ru'
    template_para_text = template.paragraphs[0].text.strip()

    for para in doc.paragraphs:  # Проверяем, есть ли уже такой параграф после target_text
        if template_para_text in para.text:
            return

    for para in doc.paragraphs:  # Ищем абзац с target_text
        if target_text in para.text:
            para_to_copy = template.paragraphs[0]
            p_xml = para_to_copy._p.xml
            e = parse_xml(p_xml)
            para._element.addnext(e)
            new_para = Paragraph(e, para._parent)
            _force_font_size(new_para, size=11)
            break


# if __name__ == '__main__':
#     mock_doc = 'mock\заявление на включение требований в РТК_2rsfdofiswdf.docx.docx'
#     # signa_path = 'templates/signa.png'
#     # temp_add_gp = 'templates/gosposhlina/add_gosposhlina.docx'
#     save_output_mock = 'mock/output.docx'
#     try:
#         doc = open_docx(mock_doc)
#         print(extract_date(doc))
#         # template = open_docx(temp_add_gp)
#         # insert_signature(doc, signa_path)
#         doc.save(save_output_mock)
#     except Exception as e:
#         print(f'Ошибка: {e}')
